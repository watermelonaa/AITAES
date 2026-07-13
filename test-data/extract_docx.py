import docx
import os
import sys

# Fix encoding for stdout
sys.stdout.reconfigure(encoding='utf-8')

os.chdir('d:/AITAES/AITAES/docs')
targets = [
    '计算机网络试题库含答案2026.docx',
    '计算机网络基础试题(十套试卷-附答案).docx'
]

for fname in targets:
    print(f'\n{"="*80}')
    print(f'FILE: {fname}')
    print(f'{"="*80}')
    if not os.path.exists(fname):
        print(f'NOT FOUND: {fname}')
        continue
    doc = docx.Document(fname)
    print(f'Paragraphs: {len(doc.paragraphs)}')
    print(f'Tables: {len(doc.tables)}')

    # Print first 200 paragraphs
    for i, p in enumerate(doc.paragraphs[:200]):
        t = p.text.strip()
        if t:
            print(f'[{i}] [{p.style.name}] {t[:300]}')

    # Print table info
    for ti, table in enumerate(doc.tables[:5]):
        print(f'\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---')
        for ri, row in enumerate(table.rows[:5]):
            cells = [cell.text.strip()[:80] for cell in row.cells]
            print(f'  Row {ri}: {cells}')
